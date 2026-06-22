from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
import os


def generate_offer_letter(candidate: dict, job_role: str, salary: str, joining_date: str) -> str:
    """
    Generate offer letter as Word document
    Input:  candidate dict, job details
    Output: path to generated .docx file
    """
    doc = Document()

    # Company Header
    header = doc.add_heading('AI HIRING ASSISTANT', 0)
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub = doc.add_paragraph('HR Department | recruitment@company.com')
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph('')

    # Date
    date_para = doc.add_paragraph(f"Date: {datetime.now().strftime('%d %B %Y')}")

    doc.add_paragraph('')

    # Candidate details
    doc.add_paragraph(f"To,")
    doc.add_paragraph(f"{candidate.get('name', 'Candidate')}")
    doc.add_paragraph(f"Email: {candidate.get('email', '')}")

    doc.add_paragraph('')

    # Subject
    subject = doc.add_paragraph()
    subject_run = subject.add_run(f"Subject: Offer Letter for the position of {job_role}")
    subject_run.bold = True

    doc.add_paragraph('')

    # Body
    doc.add_paragraph(f"Dear {candidate.get('name', 'Candidate')},")

    doc.add_paragraph(
        f"We are pleased to offer you the position of {job_role} at our organization. "
        f"After careful review of your qualifications and interview performance, "
        f"we are confident you will be a valuable addition to our team."
    )

    doc.add_paragraph('')

    # Terms
    doc.add_heading('Employment Terms:', level=2)

    terms = [
        f"Position: {job_role}",
        f"Annual CTC: {salary}",
        f"Date of Joining: {joining_date}",
        f"Employment Type: Full Time",
        f"Probation Period: 6 months"
    ]

    for term in terms:
        doc.add_paragraph(term, style='List Bullet')

    doc.add_paragraph('')

    doc.add_paragraph(
        "Please confirm your acceptance of this offer by signing and returning "
        "a copy of this letter within 7 days."
    )

    doc.add_paragraph('')
    doc.add_paragraph("We look forward to welcoming you to our team!")
    doc.add_paragraph('')
    doc.add_paragraph("Regards,")
    doc.add_paragraph("HR Manager")
    doc.add_paragraph("AI Hiring Assistant")

    # Save file
    os.makedirs('data/processed/offers', exist_ok=True)
    candidate_id = candidate.get('candidate_id', 'UNKNOWN')
    file_path = f"data/processed/offers/offer_{candidate_id}.docx"
    doc.save(file_path)

    print(f"Offer letter saved: {file_path}")
    return file_path


# ---------- TEST IT ----------
if __name__ == "__main__":
    test_candidate = {
        'candidate_id': 'CAND-001',
        'name': 'John Doe',
        'email': 'john@gmail.com'
    }

    path = generate_offer_letter(
        candidate=test_candidate,
        job_role="Data Scientist",
        salary="12 LPA",
        joining_date="01 August 2026"
    )
    print(f"Offer letter created at: {path}")